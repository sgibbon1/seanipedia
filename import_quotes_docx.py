#!/usr/bin/env python3
"""import_quotes_docx.py — import a OneNote page exported as .docx into the vault.

WHY THIS PATH EXISTS
The Graph API route (fetch_onenote_page.py) is the tidiest, but it depends on a
live token and gets throttled hard after a full export. Exporting the page from
OneNote to Word and dropping it in ~/Downloads is a reliable manual fallback that
needs no auth at all.

WHAT THE .DOCX GIVES US
Word flattens OneNote's outline: every paragraph comes through as "Normal" with
no list numbering. What survives — and what this relies on — is INDENTATION:
  - a paragraph at indent 0 starts a new entry (a quote, or a subject heading)
  - an indented paragraph continues the previous entry (wrapped/verse lines)
Subject headings are then separated from quotes the same way the indexer does it:
short, unquoted, unattributed, and followed by other entries.

Output matches the rest of the collection: YAML frontmatter, "# Title", then a
numbered list with continuation lines indented — i.e. already in the repaired
shape, so clean_quote_sources.py has nothing left to fix.

Usage:
    python3 import_quotes_docx.py --docx ~/Downloads/"By Subject.docx"
    python3 import_quotes_docx.py --docx ~/Downloads/"By Subject.docx" --apply
"""
from __future__ import annotations

import argparse
import os
import re
import shutil
from datetime import date
from pathlib import Path

import docx
from dotenv import load_dotenv

HERE = Path(__file__).parent
load_dotenv(dotenv_path=HERE / ".env", override=True)

VAULT = Path(os.environ.get("VAULT_PATH", "./vault"))
QUOTES_DIR = VAULT / "_sources" / "Τά εἰς ἑαυτόν" / "Quotes"
BACKUP = Path.home() / "scripts" / f"docx-import-backup-{date.today().isoformat()}"

_QUOTE_MARKS = '"“”«»‘’'
MAX_HEADING_CHARS = 60


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _outline_level(p) -> int | None:
    """Word's list level (ilvl) for a paragraph, or None if it isn't a list item.

    THE KEY TO THIS FORMAT. Word renders the outline markers (1. / a. / i.) as
    AUTO-NUMBERING — they are not in paragraph.text and left_indent stays 0, so
    neither text nor indentation reveals the structure. The real hierarchy lives
    in the numbering XML. In Sean's export:
        ilvl 1 = topic (1., 2., …)      ilvl 2 = quote (a., b., …)
        ilvl 3 = sub-quote (i., ii., …) None   = wrapped continuation line
    """
    npr = p._p.find(f".//{W}numPr")
    if npr is None:
        return None
    il = npr.find(f"{W}ilvl")
    return int(il.get(f"{W}val")) if il is not None else 0


def read_outline(path: Path) -> list[tuple[int, str]]:
    """Return [(level, text)] using Word's real outline levels.

    A paragraph with no list level continues the entry above it (that's how a
    two-line couplet like Harrington's arrives), so it is folded in rather than
    becoming a stray entry.
    """
    d = docx.Document(str(path))
    entries: list[list] = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        lvl = _outline_level(p)
        if lvl is None and entries:
            entries[-1][1].append(text)      # continuation of previous entry
        else:
            entries.append([lvl if lvl is not None else 1, [text]])
    return [(lvl, "\n".join(lines)) for lvl, lines in entries]


def read_html_outline(path: Path) -> list[tuple[int, str]]:
    """Read a OneNote page's exported HTML, preserving its <ol>/<li> nesting.

    The Graph API returns the page with its outline intact — real nested <ol>
    elements — which is strictly better than any copy-paste into Word (Word
    drops the structure) and better than the legacy markdownify converter, which
    flattens the nesting and re-splits multi-line quotes on <br>.

    Depth comes from counting list ancestors, so OneNote's numbering scheme
    (Arabic topic / roman quote / whatever below) maps straight onto levels
    without caring which glyph Word or the browser happens to render.
    """
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(path.read_bytes(), "html.parser")
    entries: list[list] = []
    loose_run = False

    def own_text(el) -> str:
        """Text of this element only, excluding any nested list — otherwise a
        topic would swallow every quote beneath it."""
        clone = BeautifulSoup(str(el), "html.parser")
        for nested in clone.find_all(["ol", "ul"]):
            nested.decompose()
        for br in clone.find_all("br"):
            br.replace_with("\n")
        for blk in clone.find_all(["p", "div"]):
            blk.append("\n")
        return "\n".join(l.strip() for l in clone.get_text().splitlines() if l.strip())

    # Walk <li> and <p> in DOCUMENT ORDER. Critical: OneNote emits the
    # continuation lines of a multi-line quote as <p> siblings that sit INSIDE
    # the <ol> but AFTER the closing </li> — not nested within the item. Reading
    # only <li> elements silently dropped 88 lines of poetry (St. Patrick's
    # Breastplate lost every line but its first). So a <p> whose parent is a
    # list continues the entry above it.
    for el in soup.find_all(["li", "p"]):
        if el.name == "li":
            depth = len([a for a in el.parents if a.name in ("ol", "ul")]) - 1
            text = own_text(el)
            if text:
                entries.append([max(0, depth), [text]])
                loose_run = False
        else:  # <p>
            # ORDER MATTERS. A continuation <p> sits directly inside the <ol>,
            # but that <ol> is itself inside the parent topic's <li> — so a
            # find_parent("li") test fires first and wrongly skips it. Judge by
            # the DIRECT parent: parent is a list => continuation line.
            if el.parent is not None and el.parent.name in ("ol", "ul"):
                if entries:
                    text = own_text(el)
                    if text:
                        entries[-1][1].append(text)
                    loose_run = False
            elif el.find_parent("li") is None:
                # A <p> outside every list — loose text in the page body (the
                # Dies Irae couplet trails the Faith page this way). Don't fuse
                # it into the unrelated quote above; give it its own entry, and
                # join consecutive loose lines since they form one passage.
                text = own_text(el)
                if text:
                    if loose_run and entries:
                        entries[-1][1].append(text)
                    else:
                        entries.append([0, [text]])
                        loose_run = True
            # else: inside an <li>, already captured by that item.
    return [(lvl, "\n".join(parts)) for lvl, parts in entries]


def read_blocks(path: Path) -> list[str]:
    """Legacy path: group flat paragraphs by indentation (older export style)."""
    d = docx.Document(str(path))
    blocks: list[list[str]] = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        ind = p.paragraph_format.left_indent
        indented = bool(ind and ind.inches > 0.1)
        if indented and blocks:
            blocks[-1].append(text)
        else:
            blocks.append([text])
    return ["\n".join(b) for b in blocks]


def is_heading(text: str, nxt: str | None) -> bool:
    """A subject heading ("Ideas", "Music", "War") rather than a quote.

    Conservative on purpose — same rule the indexer uses. Misfiling a heading as
    a quote costs one skim; swallowing a real quote loses it silently.
    """
    if "\n" in text or len(text) > MAX_HEADING_CHARS:
        return False
    if any(q in text for q in _QUOTE_MARKS):
        return False
    if re.search(r"[-–—]\s*[A-ZΑ-Ω]", text):     # has an attribution
        return False
    if text.rstrip().endswith((".", "!", "?", ":", ";", ",")):
        return False
    return nxt is not None                        # headings introduce something


def build_markdown_outline(entries: list[tuple[int, str]], title: str) -> tuple[str, dict]:
    """Render Word outline levels as the collection's standard nested numbering.

    Levels are normalised so the shallowest present level becomes top-level,
    then each level is numbered independently and indented 3 spaces per depth —
    matching the other files, so the indexer's numbered-mode parser reads it
    without any special-casing.
    """
    levels = sorted({lvl for lvl, _ in entries})
    base = levels[0] if levels else 1
    counters: dict[int, int] = {}
    out = [f"# {title}", ""]
    stats = {"topics": 0, "quotes": 0, "details": 0}

    for lvl, text in entries:
        depth = max(0, lvl - base)
        counters[depth] = counters.get(depth, 0) + 1
        for deeper in [k for k in list(counters) if k > depth]:
            del counters[deeper]           # restart numbering inside a new parent
        pad = "   " * depth
        lines = text.splitlines()
        out.append(f"{pad}{counters[depth]}. {lines[0]}")
        for cont in lines[1:]:
            out.append(f"{pad}   {cont}")
        stats["topics" if depth == 0 else "quotes" if depth == 1 else "details"] += 1
    return "\n".join(out) + "\n", stats


def build_markdown(blocks: list[str], title: str) -> tuple[str, int, int]:
    """Render entries as a numbered list, quotes nested under their heading."""
    out: list[str] = [f"# {title}", ""]
    n_top = n_quote = n_head = 0
    sub = 0
    have_heading = False

    for i, text in enumerate(blocks):
        nxt = blocks[i + 1] if i + 1 < len(blocks) else None
        lines = text.splitlines()
        if is_heading(text, nxt):
            n_top += 1; n_head += 1; sub = 0; have_heading = True
            out.append(f"{n_top}. {lines[0]}")
            continue
        n_quote += 1
        if have_heading:
            sub += 1
            out.append(f"   {sub}. {lines[0]}")
            for cont in lines[1:]:
                out.append(f"      {cont}")
        else:
            n_top += 1
            out.append(f"{n_top}. {lines[0]}")
            for cont in lines[1:]:
                out.append(f"   {cont}")
    return "\n".join(out) + "\n", n_quote, n_head


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", help="path to a .docx export")
    ap.add_argument("--html", help="path to a OneNote page HTML (from Graph API)")
    ap.add_argument("--title", default=None, help="defaults to the .docx filename")
    ap.add_argument("--apply", action="store_true")
    args = ap.parse_args()

    if not (args.docx or args.html):
        raise SystemExit("give --docx or --html")
    src = Path(os.path.expanduser(args.html or args.docx))
    if not src.exists():
        raise SystemExit(f"not found: {src}")
    title = args.title or src.stem

    # Prefer real structure: HTML <ol> nesting, else Word outline levels,
    # else indentation. All three feed the same renderer.
    entries = read_html_outline(src) if args.html else read_outline(src)
    if any(l is not None for l, _ in entries) and len({l for l, _ in entries}) > 1:
        body, stats = build_markdown_outline(entries, title)
        print(f"outline   : {stats['topics']} topics, {stats['quotes']} quotes, "
              f"{stats['details']} sub-details")
        n_quote, n_head = stats["quotes"], stats["topics"]
    else:
        blocks = read_blocks(src)
        body, n_quote, n_head = build_markdown(blocks, title)
        print("outline   : (no list levels found — used indentation fallback)")

    fm = (f"---\ntitle: \"{title}\"\ntype: reflection\n"
          f"notebook: \"Τά εἰς ἑαυτόν\"\nsection: \"Quotes\"\n"
          f"location: \"Τά εἰς ἑαυτόν > Quotes\"\n"
          f"source: \"imported from OneNote .docx export\"\n"
          f"imported: {date.today().isoformat()}\n---\n\n")
    new_text = fm + body

    target = QUOTES_DIR / f"{title}.md"
    print(f"docx      : {src}")
    print(f"entries   : {n_quote} quotes, {n_head} topics")
    print(f"target    : {target}")
    if target.exists():
        old = target.read_text(encoding="utf-8")
        print(f"current   : {len(old.splitlines())} lines")
        print(f"incoming  : {len(new_text.splitlines())} lines")
        for probe in ("Morrison", "Chant of Love", "Gerontion"):
            print(f"   {probe:<14} current={old.count(probe)}  incoming={body.count(probe)}")

    preview = Path("/tmp") / f"{title}_from_docx.md"
    preview.write_text(new_text, encoding="utf-8")
    print(f"preview   : {preview}")

    if not args.apply:
        print("\n[DRY RUN] vault not modified. Re-run with --apply.")
        return

    BACKUP.mkdir(parents=True, exist_ok=True)
    if target.exists():
        shutil.copy(target, BACKUP / target.name)
    target.write_text(new_text, encoding="utf-8")
    print(f"\nWrote {target}\nBackup: {BACKUP / target.name}")
    print("Next: python3 quotes_index.py")


if __name__ == "__main__":
    main()
